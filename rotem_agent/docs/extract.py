from __future__ import annotations

import hashlib
from dataclasses import dataclass, field
from pathlib import Path

from rotem_agent.docs.ocr import IMAGE_MIME, Ocr

TEXT_SUFFIXES = {".txt", ".md", ".csv"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
# A client photographing a certificate is the normal case, so an image is a
# document here rather than an unsupported file.
IMAGE_SUFFIXES = set(IMAGE_MIME)
SUPPORTED = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES | IMAGE_SUFFIXES

# Below this many characters per page, a PDF is almost certainly a scan with no
# text layer. Immigration files are full of them: certificates, apostilles and
# ministry letters arrive as photographs of paper.
MIN_CHARS_PER_PAGE = 40


@dataclass
class ExtractedDoc:
    path: Path
    text: str
    content_hash: str
    pages: int = 0
    needs_ocr: bool = False
    # Transcribed by a model rather than read from a text layer. A name or date
    # from such a document is evidence of what the machine saw, not of what the
    # certificate says, and must be confirmed against the original before any
    # discrepancy is asserted.
    machine_read: bool = False
    ocr_usage: object | None = None
    warnings: list[str] = field(default_factory=list)

    @property
    def is_empty(self) -> bool:
        return not self.text.strip()


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED


def extract_file(path: Path, ocr: Ocr | None = None) -> ExtractedDoc:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return _extract_text(path)
    if suffix in PDF_SUFFIXES:
        return _transcribe_if_needed(_extract_pdf(path), ocr)
    if suffix in DOCX_SUFFIXES:
        return _extract_docx(path)
    if suffix in IMAGE_SUFFIXES:
        return _extract_image(path, ocr)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def _transcribe_if_needed(document: ExtractedDoc, ocr: Ocr | None) -> ExtractedDoc:
    """A PDF with a text layer is never re-read: the layer is the better source."""
    if not document.needs_ocr or ocr is None:
        return document

    result = ocr.read(document.path, pages=document.pages)
    document.warnings.extend(result.warnings)
    if not result.ok:
        return document

    document.text = result.text
    document.needs_ocr = False
    document.machine_read = True
    document.ocr_usage = result.usage
    return document


def _extract_image(path: Path, ocr: Ocr | None) -> ExtractedDoc:
    digest = file_hash(path)
    if ocr is None:
        return ExtractedDoc(
            path=path,
            text="",
            content_hash=digest,
            pages=1,
            needs_ocr=True,
            warnings=["an image holds no text layer and needs OCR to be searchable"],
        )

    result = ocr.read(path, pages=1)
    return ExtractedDoc(
        path=path,
        text=result.text,
        content_hash=digest,
        pages=1,
        needs_ocr=not result.ok,
        machine_read=result.ok,
        ocr_usage=result.usage,
        warnings=list(result.warnings),
    )


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(65536), b""):
            digest.update(block)
    return digest.hexdigest()


def _extract_text(path: Path) -> ExtractedDoc:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1255", "cp1252"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")
    return ExtractedDoc(path=path, text=text.strip(), content_hash=file_hash(path))


def _extract_pdf(path: Path) -> ExtractedDoc:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf") from exc

    warnings: list[str] = []
    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        return ExtractedDoc(
            path=path,
            text="",
            content_hash=file_hash(path),
            needs_ocr=True,
            warnings=[f"could not be read ({exc})"],
        )

    text = "\n\n".join(p.strip() for p in pages if p.strip())
    page_count = len(pages)
    # A scan yields a handful of stray characters rather than nothing, so an
    # emptiness check alone would let it through as a document with no content
    # and the retrieval index would quietly lose the file.
    needs_ocr = page_count > 0 and len(text) < MIN_CHARS_PER_PAGE * page_count
    if needs_ocr:
        warnings.append(
            f"only {len(text)} characters across {page_count} page(s): "
            "this looks like a scan and needs OCR to be searchable"
        )
    return ExtractedDoc(
        path=path,
        text=text,
        content_hash=file_hash(path),
        pages=page_count,
        needs_ocr=needs_ocr,
        warnings=warnings,
    )


def _extract_docx(path: Path) -> ExtractedDoc:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx") from exc

    document = docx.Document(str(path))
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ExtractedDoc(
        path=path,
        text="\n\n".join(parts),
        content_hash=file_hash(path),
    )
